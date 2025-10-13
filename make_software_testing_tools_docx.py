#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Generate 'software_testing_tools.docx' with the provided comprehensive list,
preserving headings and bullet formatting.
Requires: python-docx  (pip install python-docx)
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text):
    t = doc.add_heading(text, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_section_heading(doc, text):
    doc.add_heading(text, level=1)

def add_tool_heading(doc, text):
    doc.add_heading(text, level=2)

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def add_spacer(doc, height_pts=6):
    p = doc.add_paragraph("")
    p_format = p.paragraph_format
    p_format.space_after = Pt(height_pts)

def build_document(doc):
    # Title
    add_title(doc, "Software Testing Tools")

    # 1. Unit Testing Tools
    add_section_heading(doc, "1. Unit Testing Tools")
    add_spacer(doc)

    add_tool_heading(doc, "JUnit")
    add_bullet(doc, "Description: JUnit is the de facto standard, open-source testing framework for the Java programming language. It provides annotations to identify test methods, assertions for testing expected results, and test runners for executing tests.")
    add_bullet(doc, "Purpose & Use Cases: Used by Java developers to write and run repeatable tests on their code. For example, testing a single method in a class to ensure it returns the correct value for a given set of inputs (assertEquals(expected, actual)). It's a core component of Test-Driven Development (TDD) in Java.")
    add_spacer(doc)

    add_tool_heading(doc, "PyTest")
    add_bullet(doc, "Description: A popular, feature-rich testing framework for Python. PyTest allows writing simple and scalable test cases, from basic unit tests to complex functional tests. It is known for its simple syntax (no boilerplate code required) and powerful features like fixtures, plugins, and detailed reporting.")
    add_bullet(doc, "Purpose & Use Cases: The go-to tool for unit testing in Python. A developer might use it to test a data processing function to ensure it handles edge cases, or to verify that a class method correctly modifies the object's state.")
    add_spacer(doc)

    add_tool_heading(doc, "Jest")
    add_bullet(doc, "Description: A delightful JavaScript Testing Framework with a focus on simplicity. Developed by Facebook, it works out-of-the-box for most JavaScript projects, especially those using React. It includes an assertion library, a test runner, and a built-in mocking library.")
    add_bullet(doc, "Purpose & Use Cases: Primarily used for testing JavaScript code, especially components in modern web frameworks like React, Angular, and Vue.js. A typical use case is to test a React component to verify that it renders correctly based on its props.")
    add_spacer(doc, 10)

    # 2. Integration Testing Tools
    add_section_heading(doc, "2. Integration Testing Tools")
    add_spacer(doc)

    add_tool_heading(doc, "Postman")
    add_bullet(doc, "Description: An API platform that allows users to design, build, test, and iterate their APIs. It has a user-friendly graphical interface for making HTTP requests (GET, POST, PUT, etc.) and inspecting responses. Its test automation capabilities allow for creating entire test suites for APIs.")
    add_bullet(doc, "Purpose & Use Cases: Ideal for testing the integration points between microservices. For example, testing if the Order Service can correctly call the Payment Service and receive a successful response, ensuring the two services are integrated correctly.")
    add_spacer(doc)

    add_tool_heading(doc, "REST Assured")
    add_bullet(doc, "Description: A Java library specifically designed for making testing REST services simple. It provides a domain-specific language (DSL) that makes it easy to write powerful, readable tests for RESTful APIs.")
    add_bullet(doc, "Purpose & Use Cases: Used by QA engineers and developers to automate integration tests for Java-based backends. A typical use case is writing a test that sends a POST request to create a user and then a GET request to validate that the user was created successfully with the correct data.")
    add_spacer(doc, 10)

    # 3. System Testing (End-to-End) Tools
    add_section_heading(doc, "3. System Testing (End-to-End) Tools")
    add_spacer(doc)

    add_tool_heading(doc, "Selenium")
    add_bullet(doc, "Description: The industry-standard, open-source framework for automating web browsers. It provides a suite of tools (WebDriver, IDE, Grid) that allows testers to write scripts in various programming languages (Java, C#, Python, etc.) to automate user actions on a web application.")
    add_bullet(doc, "Purpose & Use Cases: Used for automating the functional and regression testing of web applications. A classic use case is to automate a full user journey on an e-commerce website: logging in, searching for a product, adding it to the cart, and completing the checkout process.")
    add_spacer(doc)

    add_tool_heading(doc, "Cypress")
    add_bullet(doc, "Description: A modern, all-in-one testing framework built for the modern web. It runs directly in the browser alongside the application, providing faster and more reliable testing. It features time-travel debugging, automatic waiting, and excellent error messages.")
    add_bullet(doc, "Purpose & Use Cases: A popular alternative to Selenium for E2E testing of JavaScript-based applications. It's particularly useful for testing complex user interactions in single-page applications (SPAs) built with React, Angular, or Vue.js.")
    add_spacer(doc)

    add_tool_heading(doc, "Playwright")
    add_bullet(doc, "Description: An open-source framework developed by Microsoft for web automation and end-to-end testing. It is known for its ability to drive all modern browsers (Chromium, Firefox, WebKit) with a single API. It features auto-waits, network interception, and cross-browser emulation.")
    add_bullet(doc, "Purpose & Use Cases: Similar to Selenium and Cypress, it's used for automating E2E tests. Its strengths lie in cross-browser testing and handling modern web features. A use case could be testing a web application's responsive design by running the same test script on emulated mobile and desktop browsers.")
    add_spacer(doc, 10)

    # 4. API Testing Tools
    add_section_heading(doc, "4. API Testing Tools")
    add_spacer(doc)

    add_tool_heading(doc, "Postman (Also listed under Integration Testing)")
    add_bullet(doc, "Description: As mentioned, Postman is a comprehensive API platform. For pure API testing, its features like test scripts (using JavaScript), environment variables, collection runner, and command-line integration (Newman) are invaluable.")
    add_bullet(doc, "Purpose & Use Cases: Validating API endpoints for correct status codes, response times, and response body content. For example, testing a /users/{id} endpoint to ensure it returns a 404 Not Found error when an invalid ID is provided.")
    add_spacer(doc)

    add_tool_heading(doc, "Insomnia")
    add_bullet(doc, "Description: A powerful, open-source API client and design tool for REST, GraphQL, and gRPC. It has a clean, modern interface and provides features like environment variables, code generation, and robust plugin support.")
    add_bullet(doc, "Purpose & Use Cases: An alternative to Postman for designing, debugging, and testing APIs. It is particularly well-regarded for its excellent GraphQL support, making it ideal for testing applications that use a GraphQL API.")
    add_spacer(doc, 10)

    # 5. Load (Performance) Testing Tools
    add_section_heading(doc, "5. Load (Performance) Testing Tools")
    add_spacer(doc)

    add_tool_heading(doc, "Apache JMeter")
    add_bullet(doc, "Description: A 100% pure Java, open-source application designed to load test functional behavior and measure performance. It can be used to simulate a heavy load on a server, group of servers, network, or object to test its strength or to analyze overall performance under different load types.")
    add_bullet(doc, "Purpose & Use Cases: Used to conduct performance and stress tests on web applications, APIs, and databases. A common use case is simulating 5,000 concurrent users hitting a login page to measure server response time and identify performance bottlenecks before a product launch.")
    add_spacer(doc)

    add_tool_heading(doc, "Gatling")
    add_bullet(doc, "Description: A modern, high-performance, open-source load testing tool. Test scripts are written in a simple, expressive DSL based on Scala. Gatling is known for its excellent performance and for generating beautiful, detailed HTML reports.")
    add_bullet(doc, "Purpose & Use Cases: A developer-friendly tool for load testing. It's often used in CI/CD pipelines to continuously monitor application performance. For instance, an engineer might write a Gatling script to simulate a complex user workflow and run it after every build to catch performance regressions early.")
    add_spacer(doc)

    add_tool_heading(doc, "k6")
    add_bullet(doc, "Description: An open-source, developer-centric load testing tool built for making performance testing a productive and enjoyable experience. Test scripts are written in JavaScript (ES2015/ES6) and it focuses on setting performance goals (Service Level Objectives, or SLOs) directly in the test script.")
    add_bullet(doc, "Purpose & Use Cases: Ideal for teams that want to integrate performance testing into their development lifecycle (a \"shift-left\" approach). A typical use case is to define performance criteria, such as \"99% of API requests must complete in under 200ms,\" and fail the CI/CD build if the system doesn't meet this threshold under load.")
    add_spacer(doc, 10)

    # 6. Security Testing Tools
    add_section_heading(doc, "6. Security Testing Tools")
    add_spacer(doc)

    add_tool_heading(doc, "OWASP ZAP (Zed Attack Proxy)")
    add_bullet(doc, "Description: An open-source web application security scanner. It is one of the world's most popular free security tools and is actively maintained by a dedicated international team of volunteers. It acts as a \"man-in-the-middle proxy,\" intercepting and inspecting traffic sent between a browser and a web application to find security vulnerabilities.")
    add_bullet(doc, "Purpose & Use Cases: Used for Dynamic Application Security Testing (DAST). A security professional or developer can use ZAP to perform an automated scan of their web application to find common vulnerabilities like SQL Injection, Cross-Site Scripting (XSS), and insecure configurations.")
    add_spacer(doc)

    add_tool_heading(doc, "Burp Suite")
    add_bullet(doc, "Description: The industry-standard toolkit for web application security testing. It provides a comprehensive set of tools for every stage of the testing process, from initial mapping of an application's attack surface to finding and exploiting security vulnerabilities. It has both a free Community Edition and a more powerful Professional/Enterprise version.")
    add_bullet(doc, "Purpose & Use Cases: Used by penetration testers and security experts for in-depth security assessments. Its intercepting proxy allows for manual manipulation of requests to test for complex vulnerabilities that automated scanners might miss.")
    add_spacer(doc)

    add_tool_heading(doc, "SonarQube")
    add_bullet(doc, "Description: An open-source platform for continuous inspection of code quality. While not exclusively a security tool, its Static Application Security Testing (SAST) capabilities are a key feature. It analyzes source code to find security vulnerabilities, bugs, and \"code smells\" without actually executing the application.")
    add_bullet(doc, "Purpose & Use Cases: Integrated into the CI/CD pipeline to automatically scan code for security hotspots (e.g., hardcoded passwords, use of weak cryptographic algorithms) every time a developer commits new code, helping to prevent vulnerabilities from ever reaching production.")
    add_spacer(doc)

def main():
    doc = Document()
    # Optional: adjust page margins slightly
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    build_document(doc)
    out_name = "software_testing_tools.docx"
    doc.save(out_name)
    print(f"Created {out_name}")

if __name__ == "__main__":
    main()
